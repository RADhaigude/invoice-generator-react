import os, base64, requests, pypandoc
from openpyxl import Workbook
from docx import Document

# Token from GitHub Secrets
token = os.getenv("SONAR_TOKEN")
if not token:
    raise ValueError(" SONAR_TOKEN not set in environment!")

#  Project key (from repo name)
repo = os.getenv("GITHUB_REPO", "")
if not repo:
    raise ValueError(" GITHUB_REPO not set!")

# Example: repo = "radhaigude/invoice-generator-react"
project = repo.replace("/", "_").lower()  #  SonarCloud format

print(f" Using SonarCloud project key: {project}")

#  API call
headers = {"Authorization": "Basic " + base64.b64encode(f"{token}:".encode()).decode()}
url = f"https://sonarcloud.io/api/measures/component?component={project}&metricKeys=bugs,vulnerabilities,code_smells,coverage,duplicated_lines_density"
resp = requests.get(url, headers=headers)

if resp.status_code != 200:
    raise Exception(f" API call failed: {resp.status_code} {resp.text}")

data = resp.json()
measures = data["component"]["measures"]

# Prepare metrics
metrics = {m["metric"]: m["value"] for m in measures}
print(" Metrics fetched:", metrics)

#  Generate Markdown content
md_content = "# SonarCloud Quality Report\n\n"
for k, v in metrics.items():
    md_content += f"- **{k.capitalize()}**: {v}\n"

#  Save RTF report
pypandoc.convert_text(md_content, 'rtf', format='md',
                      outputfile='sonar_report_quality.rtf',
                      extra_args=['--standalone'])

#  Save DOCX report
doc = Document()
doc.add_heading("SonarCloud Quality Report", 0)
for k, v in metrics.items():
    doc.add_paragraph(f"{k.capitalize()}: {v}")
doc.save("sonar_report_quality.docx")

#  Save XLSX report
wb = Workbook()
ws = wb.active
ws.title = "Quality Report"
ws.append(["Metric", "Value"])
for k, v in metrics.items():
    ws.append([k, v])
wb.save("sonar_report_quality.xlsx")

print(" Reports generated successfully!")
